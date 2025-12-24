from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from exports.styles import STYLES


class SignatureSection:
    """Signature section for the contract."""

    def __init__(self, document):
        self.document = document

    def render(self) -> None:
        """Render the complete signature section."""
        self._add_heading()
        self._add_spacing(1)
        self._add_signature_table()
        self._add_spacing(1)
        self._add_date_line()

    def _add_heading(self) -> None:
        """Add signature section heading."""
        line_para = self.document.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run(STYLES.THICK_LINE)
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = STYLES.PRIMARY_COLOR

        heading_para = self.document.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading_para.add_run("✍️ AGREEMENT & SIGNATURES")
        STYLES.apply_font(heading_run, STYLES.SECTION_HEADING)

    def _add_signature_table(self) -> None:
        """Add professional signature table."""
        table = self.document.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        for col in table.columns:
            col.width = Inches(3.5)

        # Header row
        header_row = table.rows[0]
        self._style_header_cell(header_row.cells[0], "👤 CUSTOMER")
        self._style_header_cell(header_row.cells[1], "🔧 SERVICE PROVIDER")

        # Data rows
        self._add_signature_field(table.rows[1].cells[0], "Name:")
        self._add_signature_field(table.rows[1].cells[1], "Name:")
        self._add_signature_field(table.rows[2].cells[0], "Signature:")
        self._add_signature_field(table.rows[2].cells[1], "Signature:")
        self._add_signature_field(table.rows[3].cells[0], "Date:")
        self._add_signature_field(table.rows[3].cells[1], "Date:")

    def _style_header_cell(self, cell, text: str) -> None:
        """Style a header cell with background color."""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = STYLES.BORDER_COLOR

    def _add_signature_field(self, cell, label: str) -> None:
        """Add a signature field to a cell."""
        para = cell.paragraphs[0]

        label_run = para.add_run(f"{label}  ")
        STYLES.apply_font(label_run, STYLES.LABEL)

        line_run = para.add_run("_" * 25)
        line_run.font.size = Pt(11)

        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

    def _add_date_line(self) -> None:
        """Add agreement date line."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Agreement Date: ")
        STYLES.apply_font(run, STYLES.LABEL)

        line_run = para.add_run("_" * 30)
        line_run.font.size = Pt(11)

    def _add_spacing(self, lines: int = 1) -> None:
        """Add vertical spacing."""
        for _ in range(lines):
            self.document.add_paragraph()
