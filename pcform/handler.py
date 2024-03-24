import win32com.client
from jdatetime import datetime as jdatetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_date_time_to_docx(doc):
    # Create a header
    header = doc.sections[0].header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add the date and time
    now = jdatetime.now()
    # Format the date and time
    # Output: پنج‌شنبه، ۰۴ فروردین ۱۴۰۲
    formatted_date = now.strftime(
        f'%A، %d %B %Y {now.hour}:{now.minute}'
    )

    run = paragraph.add_run()
    # Change format as needed
    run.text = formatted_date 
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    #run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')



# Get the current directory of the Python script
current_directory = os.path.dirname(os.path.abspath(__file__))
# Get the parent directory (up one level)
parent_directory = os.path.dirname(current_directory)
subfolder = 'Forms'
destination_folder = os.path.join(parent_directory, subfolder)
if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)

# docx handlers
def add_user_input_to_docx(filename, data_list):
    file_name = f"{filename}.docx"
    document = Document()
    document.add_heading('Computer service', 0)
    for data in data_list:
        print("_" * 50)
        for key,value in data.items():
            if key == "_":
                data = "_" * 60
                document.add_paragraph(data)
            else:
                data = f"{key}: {value}"
                document.add_paragraph(data, style='List Number')

    add_date_time_to_docx(document)
    # its save to destination folder.
    file_path = os.path.join(destination_folder, file_name)
    document.save(file_path)


# pdf handlers
def docx_to_pdf(input_docx_file, output_pdf_file):
    # Create a new instance of the Word application
    word = win32com.client.Dispatch("Word.Application")
    try:
        # Open the Word document
        file_path = os.path.abspath(input_docx_file)
        if os.path.exists(file_path):
            doc = word.Documents.Open(file_path)

        # Save the document as PDF
        # 17 corresponds to PDF format
        doc.SaveAs(output_pdf_file, FileFormat=17)
        # Close the document
        doc.Close()

    except Exception as e:
        print(f"Error: {e}")

    finally:
        # Quit Word application
        word.Quit()
        os.remove(file_path)
        print(f"User input added to '{output_pdf_file}' successfully.")
