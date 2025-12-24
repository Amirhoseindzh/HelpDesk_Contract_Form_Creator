from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from exports.styles import STYLES

from settings.config import PDF_TERMS_AND_CONDITIONS


class TermsSection:
    """Terms and Conditions section."""

    def __init__(self, document):
        self.document = document

    def render(self) -> None:
        """Render terms and conditions."""
        self._add_spacing(2)
        self._add_heading()
        self._add_content()
        self._add_acceptance_checkbox()

    def _add_heading(self) -> None:
        """Add T&C heading."""
        line_para = self.document.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run(STYLES.THIN_LINE)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = STYLES.BORDER_COLOR

        heading_para = self.document.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading_para.add_run("📋 TERMS & CONDITIONS")
        STYLES.apply_font(heading_run, STYLES.SECTION_HEADING)
        heading_run.font.color.rgb = STYLES.PRIMARY_COLOR

    def _add_content(self) -> None:
        """Add T&C content in a styled box."""
        table = self.document.add_table(rows=1, cols=1)
        table.style = "Table Grid"

        cell = table.rows[0].cells[0]

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFEF5"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        para = cell.paragraphs[0]
        run = para.add_run(PDF_TERMS_AND_CONDITIONS)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.italic = True

        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)

    def _add_acceptance_checkbox(self) -> None:
        """Add acceptance checkbox."""
        self._add_spacing(1)

        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)

        checkbox_run = para.add_run("☐ ")
        checkbox_run.font.size = Pt(14)

        text_run = para.add_run("I have read and agree to the above Terms & Conditions")
        STYLES.apply_font(text_run, STYLES.BODY)

    def _add_spacing(self, lines: int = 1) -> None:
        """Add vertical spacing."""
        for _ in range(lines):
            self.document.add_paragraph()
