from abc import ABC, abstractmethod

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt

from ..styles import STYLES


class BaseSection(ABC):
    """Abstract base class for document sections."""

    def __init__(self, document: Document):
        self.document = document

    @property
    @abstractmethod
    def section_number(self) -> str:
        """Return section number."""
        pass

    @property
    @abstractmethod
    def section_title(self) -> str:
        """Return section title."""
        pass

    @abstractmethod
    def _add_content(self, data: dict) -> None:
        """Add section-specific content."""
        pass

    def render(self, data: dict) -> None:
        """Render the complete section."""
        self._add_heading()
        self._add_decorative_line()
        self._add_spacing(1)
        self._add_content(data)
        self._add_spacing(2)

    def _add_heading(self) -> None:
        """Add section heading."""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        num_run = para.add_run(f"{self.section_number}. ")
        STYLES.apply_font(num_run, STYLES.SECTION_HEADING)

        title_run = para.add_run(self.section_title)
        STYLES.apply_font(title_run, STYLES.SECTION_HEADING)

    def _add_decorative_line(self) -> None:
        """Add decorative line under heading."""
        para = self.document.add_paragraph()
        run = para.add_run(STYLES.THIN_LINE[:55])
        run.font.size = Pt(8)
        run.font.color.rgb = STYLES.SECONDARY_COLOR

    def _add_spacing(self, lines: int = 1) -> None:
        """Add vertical spacing."""
        for _ in range(lines):
            self.document.add_paragraph()

    def _create_info_table(self, rows_data: list) -> None:
        """Create a professional information table."""
        table = self.document.add_table(rows=len(rows_data), cols=2)
        table.style = "Table Grid"

        table.columns[0].width = STYLES.TABLE_LABEL_WIDTH
        table.columns[1].width = STYLES.TABLE_VALUE_WIDTH

        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]

            # Label cell
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            STYLES.apply_font(label_run, STYLES.LABEL)

            # Background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            label_cell._tc.get_or_add_tcPr().append(shading)

            # Value cell
            value_cell = row.cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(str(value) if value else "N/A")
            STYLES.apply_font(value_run, STYLES.BODY)

    def _add_text_box(self, text: str, prefix: str = "") -> None:
        """
        Add text in a visually distinct box.

        Moved here so ALL sections can use it!
        """
        table = self.document.add_table(rows=1, cols=1)
        table.style = "Table Grid"

        cell = table.rows[0].cells[0]

        # Background color
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        # Prefix
        if prefix:
            prefix_para = cell.paragraphs[0]
            prefix_run = prefix_para.add_run(prefix)
            prefix_run.font.bold = True
            prefix_run.font.size = Pt(10)
            prefix_run.font.color.rgb = STYLES.SECONDARY_COLOR
            text_para = cell.add_paragraph()
        else:
            text_para = cell.paragraphs[0]

        # Main text
        text_run = text_para.add_run(text if text else "N/A")
        STYLES.apply_font(text_run, STYLES.BODY)

        # Padding
        text_para.paragraph_format.space_before = Pt(6)
        text_para.paragraph_format.space_after = Pt(6)
        text_para.paragraph_format.left_indent = Inches(0.1)
