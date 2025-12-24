import os
from datetime import datetime as pydatetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from .sections.device_section import DeviceSection
from .sections.parties_section import PartiesSection
from .sections.problem_section import ProblemSection, ServiceSection
from .sections.signature_section import SignatureSection
from .sections.terms_section import TermsSection
from .styles import STYLES
from jdatetime import datetime as jdatetime

from settings.config import LOGO_PATH


class DocumentGenerator:
    """Professional document generator."""

    def __init__(self):
        self.document: Optional[Document] = None

    def generate(self, data_list: List[Dict[str, Any]], destination_path: str) -> str:
        """Generate professional DOCX document."""
        self.document = Document()

        self._setup_margins()
        self._add_header(data_list)
        self._add_logo()
        self._add_title()

        for data in data_list:
            self._add_content_sections(data)

        self._add_signature_section()
        self._add_terms_section()
        self._add_footer()

        output_path = f"{destination_path}.docx"
        self.document.save(output_path)

        return output_path

    def _setup_margins(self) -> None:
        """Configure document margins."""
        for section in self.document.sections:
            section.top_margin = STYLES.MARGINS.top
            section.bottom_margin = STYLES.MARGINS.bottom
            section.left_margin = STYLES.MARGINS.left
            section.right_margin = STYLES.MARGINS.right

    def _add_header(self, data_list: List[Dict]) -> None:
        """Add header with date/time."""
        date_str = self._extract_date(data_list)

        header = self.document.sections[0].header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        formatted_date = self._format_date(date_str)
        run = para.add_run(formatted_date)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = STYLES.SECONDARY_COLOR

    def _extract_date(self, data_list: List[Dict]) -> Optional[str]:
        """Extract date from data if available."""
        if not data_list:
            return None

        first = data_list[0]
        return first.get("created_at") or first.get("date") or first.get("createdAt")

    def _format_date(self, date_str: Optional[str]) -> str:
        """Format date for display."""
        try:
            if date_str:
                for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"]:
                    try:
                        dt = pydatetime.strptime(date_str, fmt)
                        break
                    except ValueError:
                        continue
                else:
                    dt = pydatetime.now()

                j_date = jdatetime.fromgregorian(datetime=dt)
            else:
                j_date = jdatetime.now()

            return j_date.strftime(
                f"%A، %d %B %Y - {j_date.hour:02d}:{j_date.minute:02d}"
            )

        except Exception:
            return jdatetime.now().strftime("%Y/%m/%d")

    def _add_logo(self) -> None:
        """Add company logo if available."""
        if not LOGO_PATH or not os.path.exists(LOGO_PATH):
            return

        try:
            logo_para = self.document.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(LOGO_PATH, width=Inches(1.8))
            self._add_spacing(1)
        except Exception as e:
            print(f"Warning: Could not add logo: {e}")

    def _add_title(self) -> None:
        """Add document title section."""
        self._add_spacing(1)

        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("COMPUTER SERVICE CONTRACT")
        STYLES.apply_font(title_run, STYLES.TITLE)
        title_run.font.color.rgb = STYLES.PRIMARY_COLOR

        subtitle_para = self.document.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("Service Agreement & Work Order")
        STYLES.apply_font(subtitle_run, STYLES.SUBTITLE)

        line_para = self.document.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run(STYLES.THICK_LINE)
        line_run.font.size = Pt(12)
        line_run.font.color.rgb = STYLES.SECONDARY_COLOR

        self._add_spacing(2)

    def _add_content_sections(self, data: Dict[str, Any]) -> None:
        """Add all content sections."""
        PartiesSection(self.document).render(data)
        DeviceSection(self.document).render(data)
        ProblemSection(self.document).render(data)
        ServiceSection(self.document).render(data)

    def _add_signature_section(self) -> None:
        """Add signature section."""
        SignatureSection(self.document).render()

    def _add_terms_section(self) -> None:
        """Add terms and conditions."""
        TermsSection(self.document).render()

    def _add_footer(self) -> None:
        """Add document footer."""
        self._add_spacing(2)

        line_para = self.document.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run(STYLES.THICK_LINE)
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = STYLES.BORDER_COLOR

        footer_para = self.document.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            "This document serves as an official service contract and work order.\n"
            "Thank you for your business! 🙏"
        )
        STYLES.apply_font(footer_run, STYLES.FOOTER)

        doc_id_para = self.document.add_paragraph()
        doc_id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_id_run = doc_id_para.add_run(
            f"Document ID: {pydatetime.now().strftime('%Y%m%d%H%M%S')}"
        )
        doc_id_run.font.size = Pt(8)
        doc_id_run.font.color.rgb = STYLES.BORDER_COLOR

    def _add_spacing(self, lines: int = 1) -> None:
        """Add vertical spacing."""
        for _ in range(lines):
            self.document.add_paragraph()


def form_saveto_docx_handler(
    data_list: List[Dict[str, Any]], destination_folder: str
) -> str:
    """Generate DOCX file - drop-in replacement."""
    generator = DocumentGenerator()
    return generator.generate(data_list, destination_folder)
