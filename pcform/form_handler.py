import win32com.client
from jdatetime import datetime as jdatetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def form_saveto_docx_handler(data_list, destination_folder):    
    document = Document()
    document.add_heading('Computer service', 0)
    for data in data_list:
        for key,value in data.items():
            if key == "_":
                data = "_" * 60
                document.add_paragraph(data)
            else:
                data = f"{key}: {value}"
                document.add_paragraph(data, style='List Number')

        add_date_time_to_docx(document)
        # its save to destination folder.
        document.save(f"{destination_folder}.docx")


def add_date_time_to_docx(doc):
    # Create a header
    header = doc.sections[0].header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add the date and time
    now = jdatetime.now()
    # Format the date and time
    # Output: پنج‌شنبه، ۰۴ فروردین ۱۴۰۲
    formatted_date = now.strftime(
        f'%A، %d %B %Y {now.hour}:{now.minute}'
    )

    run = paragraph.add_run()
    # Change format as needed
    run.text = formatted_date 
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    #run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

def convert_docx_to_pdf(docx_file_path, pdf_file_path):
    """
    Convert a DOCX file to PDF.
    """
    word = win32com.client.Dispatch("Word.Application")
    doc = None  # Initialize doc variable
    try:
        # Open the Word document
    
        doc = word.Documents.Open(docx_file_path)

        # Save the document as PDF
        pdf_file_format = 17  # PDF format
        doc.SaveAs(pdf_file_path, FileFormat=pdf_file_format)

    except Exception as e:
        print(f"Error: {e}")

    finally:
        if doc is not None:
            # Close the document
            doc.Close()
        # Quit Word application
        word.Quit()

def form_docx_to_pdf_handler(data_list, destination_folder):
    # Generate a DOCX file with provided data
    form_saveto_docx_handler(data_list, destination_folder)
    
    docx_file_path = f"{destination_folder}.docx"
    docx_file_path = os.path.abspath(docx_file_path)

    if os.path.exists(docx_file_path):
        pdf_file_path = f"{destination_folder}.pdf"
        convert_docx_to_pdf(docx_file_path, pdf_file_path)
        os.remove(docx_file_path)
        print(f"DOCX file converted to PDF: {pdf_file_path}")

    elif not docx_file_path:
        print("Error: DOCX file path is invalid or empty.")
        
    # Convert the DOCX file to PDF
    
    
